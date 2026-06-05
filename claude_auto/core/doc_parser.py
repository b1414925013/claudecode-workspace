"""文档解析模块 - 支持 .md / .doc / .docx / .pdf 格式解析"""

from pathlib import Path
from typing import Any, Dict, List

import logging

logger = logging.getLogger(__name__)


class DocParser:
    """多格式文档解析器，将各类文档提取为纯文本"""

    SUPPORTED_EXTENSIONS = {".md", ".doc", ".docx", ".pdf"}

    def __init__(self, docs_path: Path):
        self.docs_path = Path(docs_path)

    # ------------------------------------------------------------------
    # 公共方法
    # ------------------------------------------------------------------

    def scan(self) -> List[Path]:
        """扫描文档目录，返回所有支持格式的文件列表（按文件名排序）"""
        if not self.docs_path.exists():
            raise FileNotFoundError(f"文档目录不存在: {self.docs_path}")

        files: List[Path] = []
        for ext in self.SUPPORTED_EXTENSIONS:
            files.extend(self.docs_path.glob(f"*{ext}"))
            files.extend(self.docs_path.glob(f"*{ext.upper()}"))

        # 按文件名去重排序
        seen: set = set()
        unique: List[Path] = []
        for f in sorted(files, key=lambda x: x.name):
            if f.name not in seen:
                seen.add(f.name)
                unique.append(f)
        return unique

    def parse(self, file_path: Path) -> Dict[str, Any]:
        """解析单个文档

        Returns:
            {"filename": str, "content": str}
        """
        ext = file_path.suffix.lower()
        parser = {
            ".md": self._parse_markdown,
            ".doc": self._parse_word,
            ".docx": self._parse_word,
            ".pdf": self._parse_pdf,
        }.get(ext)

        if parser is None:
            raise ValueError(f"不支持的文件格式: {ext}")

        content = parser(file_path)
        return {"filename": file_path.name, "content": content}

    def parse_all(self) -> List[Dict[str, Any]]:
        """解析文档目录下所有支持的文件

        Returns:
            解析结果列表，解析失败的文档会被记录日志但不会中断流程
        """
        documents: List[Dict[str, Any]] = []
        for file_path in self.scan():
            try:
                documents.append(self.parse(file_path))
                logger.info("[OK] 解析成功: %s", file_path.name)
            except Exception as e:
                logger.error("[FAIL] 解析失败 %s: %s", file_path.name, e)
        return documents

    # ------------------------------------------------------------------
    # 格式特定解析方法
    # ------------------------------------------------------------------

    @staticmethod
    def _parse_markdown(path: Path) -> str:
        with open(path, "r", encoding="utf-8") as f:
            return f.read().strip()

    @staticmethod
    def _parse_word(path: Path) -> str:
        from docx import Document

        doc = Document(str(path))

        # 提取段落文本
        parts = [p.text for p in doc.paragraphs if p.text.strip()]

        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))

        return "\n".join(parts)

    @staticmethod
    def _parse_pdf(path: Path) -> str:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        pages: List[str] = []
        for page in reader.pages:
            text = page.extract_text()
            if text and text.strip():
                pages.append(text.strip())
        return "\n\n".join(pages)
