"""测试用例生成与执行模块"""

import json
import logging
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List

from core.claude_client import ClaudeClient

logger = logging.getLogger(__name__)


def execute_test_cases(
    client: ClaudeClient,
    test_cases: List[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    """逐条调用 Claude API 执行测试用例"""
    results: List[Dict[str, Any]] = []
    total = len(test_cases)

    for i, tc in enumerate(test_cases, 1):
        case_id = tc.get("case_id", f"TC-{i}")
        logger.info("执行测试 [%d/%d]: %s", i, total, case_id)
        try:
            result = client.execute_test_case(tc)
            results.append(result)
            logger.info("  → %s: %s", case_id, result.get("status", "unknown"))
        except Exception as e:
            logger.error("  → %s 执行异常: %s", case_id, e)
            results.append({
                "case_id": case_id,
                "status": "error",
                "actual_result": "",
                "error_info": str(e),
                "execution_log": [],
            })

    return results


class TestCaseGenerator:
    """测试用例生成器：调用 Claude API 从文档提取测试用例"""

    def __init__(self, client: ClaudeClient, output_dir: Path):
        self.client = client
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate(self, documents: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """为所有文档生成测试用例

        Args:
            documents: DocParser.parse_all() 的输出

        Returns:
            合并后的测试用例列表
        """
        all_cases: List[Dict[str, Any]] = []
        for doc in documents:
            try:
                logger.info("正在为 [%s] 生成测试用例...", doc["filename"])
                cases = self.client.generate_test_cases(doc["content"], doc["filename"])
                all_cases.extend(cases)
                logger.info("  → 生成 %d 条测试用例", len(cases))
            except Exception as e:
                logger.error("生成失败 [%s]: %s", doc["filename"], e, exc_info=True)
        return all_cases

    def save(self, test_cases: List[Dict[str, Any]]) -> Path | None:
        """将测试用例保存为 JSON + DOCX + XLSX 三格式

        Returns:
            JSON 文件路径，若无用例则返回 None
        """
        if not test_cases:
            logger.warning("没有测试用例可保存")
            return None

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        # --- JSON ---
        data = {
            "generated_at": datetime.now().isoformat(),
            "total": len(test_cases),
            "test_cases": test_cases,
        }
        json_path = self.output_dir / f"test_cases_{timestamp}.json"
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        logger.info("JSON 已保存: %s", json_path)

        # --- DOCX ---
        try:
            docx_path = self._save_as_docx(test_cases, timestamp)
            logger.info("DOCX 已保存: %s", docx_path)
        except Exception as e:
            logger.warning("DOCX 生成失败: %s", e)

        # --- XLSX ---
        try:
            xlsx_path = self._save_as_xlsx(test_cases, timestamp)
            logger.info("XLSX 已保存: %s", xlsx_path)
        except Exception as e:
            logger.warning("XLSX 生成失败: %s", e)

        logger.info("测试用例已保存 (%d 条)", len(test_cases))
        return json_path

    # ------------------------------------------------------------------
    # DOCX 输出
    # ------------------------------------------------------------------

    def _save_as_docx(self, test_cases: List[Dict[str, Any]], timestamp: str) -> Path:
        """将测试用例输出为 Word 文档"""
        from docx import Document
        from docx.shared import Pt, Inches, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn

        doc = Document()

        # 页边距
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)

        # ── 标题 ──
        title = doc.add_heading("测试用例", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"用例总数：{len(test_cases)}")
        doc.add_paragraph("")  # 空行

        # ── 表格 ──
        headers = ["用例编号", "模块", "功能点", "前置条件", "操作步骤", "预期结果", "优先级"]
        table = doc.add_table(rows=1 + len(test_cases), cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 表头
        header_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            header_cells[i].text = h
            for p in header_cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            # 背景色
            shading = header_cells[i]._tc.get_or_add_tcPr()
            shading_elm = shading.makeelement(qn("w:shd"), {
                qn("w:val"): "clear",
                qn("w:color"): "auto",
                qn("w:fill"): "2B579A",
            })
            shading.append(shading_elm)

        # 数据行
        for row_idx, tc in enumerate(test_cases):
            row = table.rows[row_idx + 1]
            steps_text = "\n".join(f"{i+1}. {s}" for i, s in enumerate(tc.get("steps", []))) if tc.get("steps") else "—"

            values = [
                tc.get("case_id", ""),
                tc.get("module", ""),
                tc.get("feature", ""),
                tc.get("precondition", "") or "—",
                steps_text,
                tc.get("expected", ""),
                tc.get("priority", ""),
            ]
            for col_idx, val in enumerate(values):
                row.cells[col_idx].text = val
                for p in row.cells[col_idx].paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)

        # 列宽
        widths = [Cm(2.8), Cm(2.2), Cm(2.5), Cm(3.0), Cm(5.0), Cm(3.5), Cm(1.5)]
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width

        filepath = self.output_dir / f"test_cases_{timestamp}.docx"
        doc.save(str(filepath))
        return filepath

    # ------------------------------------------------------------------
    # XLSX 输出
    # ------------------------------------------------------------------

    def _save_as_xlsx(self, test_cases: List[Dict[str, Any]], timestamp: str) -> Path:
        """将测试用例输出为 Excel 工作簿"""
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter

        wb = Workbook()
        ws = wb.active
        ws.title = "测试用例"

        headers = ["用例编号", "模块", "功能点", "前置条件", "操作步骤", "预期结果", "优先级"]
        col_widths = [18, 14, 16, 20, 40, 24, 8]

        # ── 样式 ──
        header_font = Font(name="微软雅黑", bold=True, color="FFFFFF", size=11)
        header_fill = PatternFill(start_color="2B579A", end_color="2B579A", fill_type="solid")
        header_align = Alignment(horizontal="center", vertical="center")
        cell_font = Font(name="微软雅黑", size=10)
        cell_align = Alignment(vertical="top", wrap_text=True)
        thin_border = Border(
            left=Side(style="thin", color="D0D0D0"),
            right=Side(style="thin", color="D0D0D0"),
            top=Side(style="thin", color="D0D0D0"),
            bottom=Side(style="thin", color="D0D0D0"),
        )

        # ── 表头 ──
        for col, (header, width) in enumerate(zip(headers, col_widths), 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_align
            cell.border = thin_border
            ws.column_dimensions[get_column_letter(col)].width = width

        # ── 数据 ──
        for row_idx, tc in enumerate(test_cases, 2):
            steps_text = "\n".join(f"{i+1}. {s}" for i, s in enumerate(tc.get("steps", []))) if tc.get("steps") else "—"
            values = [
                tc.get("case_id", ""),
                tc.get("module", ""),
                tc.get("feature", ""),
                tc.get("precondition", "") or "—",
                steps_text,
                tc.get("expected", ""),
                tc.get("priority", ""),
            ]
            for col, val in enumerate(values, 1):
                cell = ws.cell(row=row_idx, column=col, value=val)
                cell.font = cell_font
                cell.alignment = cell_align
                cell.border = thin_border

        # ── 冻结首行 ──
        ws.freeze_panes = "A2"

        # ── 自动筛选 ──
        ws.auto_filter.ref = f"A1:G{len(test_cases) + 1}"

        # ── 行高自适应 ──
        for row_idx in range(2, len(test_cases) + 2):
            max_lines = 1
            for col in range(1, 8):
                cell_val = ws.cell(row=row_idx, column=col).value
                if cell_val:
                    max_lines = max(max_lines, str(cell_val).count("\n") + 1)
            ws.row_dimensions[row_idx].height = max(20, max_lines * 16)

        filepath = self.output_dir / f"test_cases_{timestamp}.xlsx"
        wb.save(str(filepath))
        return filepath
